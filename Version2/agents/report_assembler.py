import os
import subprocess
from docx import Document

class ReportAssemblerAgent:
    def save_and_convert(self, prs, export_format="odp"):
        pptx_file = "one_column_eda_report.pptx"
        output_file = f"one_column_eda_report.{export_format}"
        prs.save(pptx_file)
        try:
            if export_format == "odp":
                subprocess.run(["libreoffice", "--headless", "--convert-to", "odp", pptx_file], check=True)
            elif export_format == "pdf":
                subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", pptx_file], check=True)
            elif export_format == "docx":
                doc = Document()
                for slide in prs.slides:
                    for shape in slide.shapes:
                        if shape.has_text_frame:
                            doc.add_paragraph(shape.text_frame.text)
                        elif shape.shape_type == 13:  # Picture
                            doc.add_paragraph(f"[Image: {shape.name}]")
                    doc.add_page_break()
                doc.save(output_file)
                return True, open(output_file, "rb").read()
            with open(output_file, "rb") as f:
                output_bytes = f.read()
            os.remove(pptx_file)
            os.remove(output_file)
            return True, output_bytes
        except Exception as e:
            return False, f"Error converting to {export_format}: {str(e)}"

    def assemble_report(self, csv_file, col, plot_type, min_slides, user_prompt, theme, font_style, data_loader, content_gen, slide_builder, plot_gen, edited_slides=None):
        success, message = data_loader.load_data(csv_file)
        if not success:
            return False, message
        data_loader.set_column(col)
        
        slide_builder.set_theme(theme)
        slide_builder.set_font_style(font_style)
        
        # Title slide
        title_prompt = f"Analyze CSV: Rows={len(data_loader.df)}, Cols={data_loader.num_cols}, Selected={col}. Generate a 5-word title based on data and '{user_prompt}'."
        cover_title = content_gen.generate_content(title_prompt).split('\n')[0]
        slide_builder.add_title_slide(cover_title)
        
        # Overview slide(s)
        slide_titles = ["Overview of Upcoming Slides", "Introduction to Analysis"]
        slide_titles.extend([f"Comparison Plot: {col} vs {other_col}" for other_col in data_loader.other_cols])
        slide_titles.extend([f"Comparison Insights: {col} vs {other_col}" for other_col in data_loader.other_cols])
        slide_titles.extend([f"Detailed Insights: {col} vs {other_col}" for other_col in data_loader.other_cols])
        extra_slides_needed = min_slides > (2 * data_loader.num_cols)
        if extra_slides_needed:
            slide_titles.append("Index of Slides")
        if user_prompt.lower() != "default analysis of one column vs others" and "summary" in user_prompt.lower():
            slide_titles.append("Summary of Findings")
        slide_titles.append("Conclusion of Analysis")
        overview_content = [f"{i + 1}. {title}" for i, title in enumerate(slide_titles[2:-1])]
        max_points_per_slide = 6
        for i in range(0, len(overview_content), max_points_per_slide):
            chunk = overview_content[i:i + max_points_per_slide]
            title = "Overview of Upcoming Slides" if i == 0 else "Overview of Upcoming Slides Continued"
            slide_builder.add_slide(title, chunk)
        
        # Introduction slide with CSV analysis
        stats_summary = "\n".join([f"{col}: {', '.join([f'{k}={v}' for k, v in stats.items()])}" for col, stats in data_loader.stats.items()])
        intro_prompt = f"Introduce analysis of {col} vs others based on CSV with {len(data_loader.df)} rows, {data_loader.num_cols} columns, focusing on {col}. Use this analysis: '{stats_summary}' in 5 to 6 bullet points based on '{user_prompt}'."
        intro_text = content_gen.generate_content(intro_prompt)
        intro_points = content_gen.split_into_bullets(intro_text)
        slide_builder.add_slide("Introduction to Analysis", intro_points)
        
        # Comparison slides
        for other_col in data_loader.other_cols:
            chart_path, actual_plot_type = plot_gen.generate_plot(data_loader.df, col, other_col, plot_type)
            slide_builder.add_slide(f"Comparison Plot: {col} vs {other_col}", chart_path=chart_path)
            
            corr = data_loader.stats[col].get(f"corr_with_{other_col}", "N/A")
            stats_content = f"{col} vs {other_col}: Corr={corr}, {col} {list(data_loader.stats[col].items())[:3]}, {other_col} {list(data_loader.stats[other_col].items())[:3]}"
            content_points = [
                f"Rows analyzed: {len(data_loader.df)}. Total entries in CSV.",
                f"Correlation: {corr}. Shows {col} vs {other_col} link." if corr != "N/A" else f"{col} type: {data_loader.data_types[col]}. Non-numeric data detected.",
                f"{other_col} mean: {data_loader.stats[other_col]['mean']}. Average from CSV data." if 'mean' in data_loader.stats[other_col] else f"{other_col} unique: {data_loader.stats[other_col]['unique']}. Distinct values counted.",
                f"{other_col} min: {data_loader.stats[other_col]['min']}. Minimum value in CSV." if 'min' in data_loader.stats[other_col] else f"{other_col} top: {data_loader.stats[other_col]['top']}. Most frequent in CSV.",
                f"{other_col} max: {data_loader.stats[other_col]['max']}. Maximum value in CSV." if 'max' in data_loader.stats[other_col] else f"{other_col} diversity: {'High' if int(data_loader.stats[other_col]['unique']) > 5 else 'Low'}. Variation in data.",
                f"{col} stat: {data_loader.stats[col]['mean']}. Numeric average from CSV." if 'mean' in data_loader.stats[col] else f"{col} top: {data_loader.stats[col]['top']}. Top category in CSV."
            ]
            slide_builder.add_slide(f"Comparison Insights: {col} vs {other_col}", content_points, layout="text")
            
            detail_prompt = f"Provide detailed insights for {col} vs {other_col} based on CSV data: '{stats_content}', in 5 to 6 bullet points based on '{user_prompt}'."
            detail_text = content_gen.generate_content(detail_prompt)
            detail_points = content_gen.split_into_bullets(detail_text)
            slide_builder.add_slide(f"Detailed Insights: {col} vs {other_col}", detail_points)
            
            os.remove(chart_path)
        
        # Index slide
        if extra_slides_needed:
            index_content = [f"{i + 1}. {title}" for i, title in enumerate(slide_titles[2:-1])]
            index_points = content_gen.split_into_bullets("\n".join(index_content))
            slide_builder.add_slide("Index of Slides", index_points)
        
        # Summary slide
        if user_prompt and user_prompt.lower() != "default analysis of one column vs others" and "summary" in user_prompt.lower():
            summary_prompt = f"Summarize analysis of {col} vs others based on CSV data with {len(data_loader.df)} rows, {data_loader.num_cols} columns, using this analysis: '{stats_summary}' in 5 to 6 bullet points based on '{user_prompt}'."
            summary_text = content_gen.generate_content(summary_prompt)
            summary_points = content_gen.split_into_bullets(summary_text)
            slide_builder.add_slide("Summary of Findings", summary_points)
            slide_titles.append("Summary of Findings")
        
        # Additional slides
        current_slides = len(slide_titles) + 1
        if current_slides < min_slides:
            for i in range(min_slides - current_slides):
                extra_prompt = f"Provide extra analysis for {col} vs others based on CSV data with {len(data_loader.df)} rows, {data_loader.num_cols} columns, using this analysis: '{stats_summary}' in 5 to 6 bullet points based on '{user_prompt}'."
                extra_text = content_gen.generate_content(extra_prompt)
                extra_points = content_gen.split_into_bullets(extra_text)
                slide_builder.add_slide(f"Additional Analysis {i + 1}", extra_points, progress=(i + 1) / (min_slides - current_slides + 1), layout="progress")
                slide_titles.append(f"Additional Analysis {i + 1}")
        
        # Conclusion slide with CSV analysis
        conclusion_prompt = f"Conclude analysis of {col} vs others based on CSV data with {len(data_loader.df)} rows, {data_loader.num_cols} columns, using this analysis: '{stats_summary}' in 5 to 6 bullet points based on '{user_prompt}'."
        conclusion_text = content_gen.generate_content(conclusion_prompt)
        conclusion_points = content_gen.split_into_bullets(conclusion_text)
        slide_builder.add_slide("Conclusion of Analysis", conclusion_points)
        
        # Thank You slide
        slide_builder.add_title_slide("Thank You")
        
        if edited_slides:
            for slide_title, slide_content in edited_slides.items():
                slide_builder.add_slide(slide_title, slide_content)
        
        return True, slide_titles[1:]

    def finalize_report(self, export_format):
        return self.save_and_convert(self.prs, export_format)